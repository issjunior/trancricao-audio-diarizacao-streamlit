import os
import traceback
import streamlit as st
from dotenv import load_dotenv
import whisper
from pyannote.audio import Pipeline
from docx import Document
import pandas as pd
from io import BytesIO
import queue
import multiprocessing
import psutil
from datetime import timedelta

# 🖥️ Configuração da página
st.set_page_config(layout="wide", page_title="SPAV - Transcrição", page_icon="🎙️")
os.environ["SPEECHBRAIN_LOCAL_CACHE_STRATEGY"] = "copy"

# 💻 Configurações de otimização
num_cores = multiprocessing.cpu_count()
mem_total_gb = psutil.virtual_memory().total / (1024 ** 3)
disk = psutil.disk_usage("/")
total_gb = disk.total / (1024 ** 3)
used_gb = disk.used / (1024 ** 3)
free_gb = disk.free / (1024 ** 3)

# Usa sempre todos os núcleos disponíveis
threads = num_cores

# Aplica as configurações de threads
os.environ["OMP_NUM_THREADS"] = str(threads)
os.environ["MKL_NUM_THREADS"] = str(threads)
os.environ["NUMEXPR_NUM_THREADS"] = str(threads)

# 🔧 Inicialização das configurações de sessão
if 'habilitar_diarizacao' not in st.session_state:
    st.session_state['habilitar_diarizacao'] = True
if 'chunk_processing' not in st.session_state:
    st.session_state['chunk_processing'] = True
if 'auto_cleanup' not in st.session_state:
    st.session_state['auto_cleanup'] = True

# 🕰️ Função para criar SRT
def criar_srt(falas):
    srt_content = []
    for i, fala in enumerate(falas, 1):
        minutos_inicio, segundos_inicio = map(int, fala["tempo"].split(" - ")[0].split(":"))
        minutos_fim, segundos_fim = map(int, fala["tempo"].split(" - ")[1].split(":"))
        
        tempo_inicio = timedelta(minutes=minutos_inicio, seconds=segundos_inicio)
        tempo_fim = timedelta(minutes=minutos_fim, seconds=segundos_fim)
        
        srt_timestamp_inicio = str(tempo_inicio).replace(".", ",")[:12]
        srt_timestamp_fim = str(tempo_fim).replace(".", ",")[:12]
        
        srt_texto = f"{fala['locutor']}: {fala['texto'].strip('\"')}"
        
        srt_entry = (
            f"{i}\n"
            f"{srt_timestamp_inicio} --> {srt_timestamp_fim}\n"
            f"{srt_texto}\n\n"
        )
        srt_content.append(srt_entry)
    
    return "".join(srt_content)

# 🛠️ Funções auxiliares
def formatar_tempo(tempo_em_segundos):
    minutos = int(tempo_em_segundos // 60)
    segundos = int(tempo_em_segundos % 60)
    return f"{minutos:02}:{segundos:02}"

def atualizar_progresso(progresso_bar, status_text, etapa, valor, detalhes=""):
    status_msg = f"{etapa}"
    if detalhes:
        status_msg += f" - {detalhes}"
    status_text.text(status_msg)
    progresso_bar.progress(int(min(valor, 100)))

def processar_audio_chunk(modelo, audio_path, language, progress_queue):
    try:
        lang_param = None if language == "auto" else language
        resultado = modelo.transcribe(
            audio_path,
            language=lang_param,
            verbose=False,
            fp16=False,
            temperature=0.0
        )
        progress_queue.put(("transcricao_completa", resultado))
    except Exception as e:
        progress_queue.put(("erro", str(e)))

def processar_diarizacao(audio_path, token, progress_queue):
    try:
        pipeline = Pipeline.from_pretrained(
            "pyannote/speaker-diarization",
            use_auth_token=token
        )
        diarization = pipeline(audio_path)
        progress_queue.put(("diarizacao_completa", diarization))
    except Exception as e:
        progress_queue.put(("erro", str(e)))

# 🎛️ Configuração da barra lateral
opcoes_modelos = {
    "tiny": {"nome": "Tiny", "descricao": "Ultra rápido, baixa precisão", "tamanho": "39MB"},
    "base": {"nome": "Base", "descricao": "Rápido, precisão moderada", "tamanho": "74MB"},
    "small": {"nome": "Small", "descricao": "Balanceado", "tamanho": "244MB"},
    "medium": {"nome": "Medium", "descricao": "Mais lento, boa precisão", "tamanho": "769MB"},
    "large": {"nome": "Large", "descricao": "Muito lento, alta precisão", "tamanho": "1.5GB"},
    "large-v2": {"nome": "Large-v2", "descricao": "Versão 2 do modelo Large, melhor precisão", "tamanho": "1.9GB"},
    "large-v3": {"nome": "Large-v3", "descricao": "Versão 3 do modelo Large, precisão avançada", "tamanho": "2.1GB"},
    "turbo": {"nome": "Turbo", "descricao": "Modelo otimizado de alta performance", "tamanho": "2.3GB"}
}

modelo_escolhido = st.sidebar.selectbox(
    "Modelo:",
    options=list(opcoes_modelos.keys()),
    index=2,  # Corresponde ao "small"
    format_func=lambda x: opcoes_modelos[x]["nome"],
    disabled="audio_processado" in st.session_state
)

st.sidebar.text(f"Tamanho do modelo escolhido: {opcoes_modelos[modelo_escolhido]['tamanho']}")

# 🌐 Idioma
st.sidebar.subheader("🌍 Idioma do áudio")
opcoes_idiomas = {
    "pt": "Português (Brasil)",
    "en": "Inglês",
    "es": "Espanhol",
    "fr": "Francês",
    "de": "Alemão",
    "auto": "🌐 Detectar automaticamente"
}
idioma_escolhido_label = st.sidebar.selectbox(
    "Idioma:", list(opcoes_idiomas.values()), index=0,
    disabled="audio_processado" in st.session_state
)
idioma_escolhido = list(opcoes_idiomas.keys())[list(opcoes_idiomas.values()).index(idioma_escolhido_label)]

# ⚙️ Opções Avançadas
st.sidebar.subheader("⚙️ Opções Avançadas")
if "audio_processado" not in st.session_state:
    st.session_state['habilitar_diarizacao'] = st.sidebar.checkbox(
        "👥 Reconhecimento de locutor (Diarização)", 
        value=st.session_state['habilitar_diarizacao']
    )
    st.session_state['chunk_processing'] = st.sidebar.checkbox(
        "🧩 Processamento em chunks", 
        value=st.session_state['chunk_processing']
    )
    st.session_state['auto_cleanup'] = st.sidebar.checkbox(
        "🧹 Limpeza automática de memória", 
        value=st.session_state['auto_cleanup']
    )
else:
    st.sidebar.write("👥 Reconhecimento de locutor: " + 
                     (f"<span style='color:green'>Habilitado</span>" if st.session_state['habilitar_diarizacao'] else 
                      f"<span style='color:red'>Desabilitado</span>"), 
                     unsafe_allow_html=True)
    st.sidebar.write("🧩 Processamento em chunks: " + 
                     (f"<span style='color:green'>Habilitado</span>" if st.session_state['chunk_processing'] else 
                      f"<span style='color:red'>Desabilitado</span>"), 
                     unsafe_allow_html=True)
    st.sidebar.write("🧹 Limpeza automática de memória: " + 
                     (f"<span style='color:green'>Habilitada</span>" if st.session_state['auto_cleanup'] else 
                      f"<span style='color:red'>Desabilitada</span>"), 
                     unsafe_allow_html=True)

st.sidebar.markdown(
    f"💻 **Detecção automática de hardware**\n\n"
    f"- Threads de processamento: {threads}\n"
    f"- Memória RAM: {mem_total_gb:.1f} GB\n"
    f"- Espaço em disco livre: {free_gb:.1f} GB de {total_gb:.1f} GB"
)

# 🎙️ Página principal
st.header("🎙️ SPAV - Transcrição de Áudio", divider=True)
st.write("Transcrição e reconhecimento de voz.")
load_dotenv()
HUGGINGFACE_TOKEN = os.getenv("HUGGINGFACE_TOKEN")
if st.session_state['habilitar_diarizacao'] and not HUGGINGFACE_TOKEN:
    st.error("⚠️ Token do HuggingFace não encontrado. Defina no arquivo `.env` ou desabilite a diarização")
    st.stop()

# 📤 Upload do áudio
st.header("📤 Upload do Arquivo")
audio_file = st.file_uploader(
    "Selecione um arquivo de áudio", type=["mp3", "wav", "m4a", "flac"],
    help="Formatos suportados: MP3, WAV, M4A, FLAC"
)

# 🔄 Botão de reiniciar sempre visível
col1, col2, col3 = st.columns(3)
with col1:
    if st.button("🔄 Reiniciar Aplicação", type="primary"):
        for key in list(st.session_state.keys()):
            del st.session_state[key]
        st.rerun()
with col2:
    st.write("")  # Espaço reservado
with col3:
    st.write("")  # Espaço reservado

if audio_file:
    try:
        file_size_mb = len(audio_file.read()) / (1024 * 1024)
        audio_file.seek(0)
        if file_size_mb > 100:
            st.warning("⚠️ Arquivo é grande. O processamento será lento.")
    except Exception:
        st.info("✅ Arquivo carregado com sucesso")
    if "audio_processado" not in st.session_state or st.session_state["audio_processado"] != audio_file.name:
        for key in ["tabela_falas", "nome_base", "modelo_escolhido", "doc_word", "csv_data", "srt_data"]:
            st.session_state.pop(key, None)
        st.session_state["audio_processado"] = audio_file.name

# 🎬 Processamento do áudio
if audio_file:
    # Remova as colunas anteriores
    if st.button("🚀 Iniciar Transcrição", type="secondary"):
        audio_path = os.path.join(os.getcwd(), f"temp_{audio_file.name}")
        try:
            with open(audio_path, "wb") as f:
                f.write(audio_file.read())
            
            # Use st.empty() com uma largura total
            st.header("🔄 Progresso da Transcrição")
            
            # Crie um contêiner de largura total
            progress_container = st.container()
            with progress_container:
                # Use st.empty() de forma semelhante ao seu código original
                progresso_placeholder = st.empty()
                status_placeholder = st.empty()
                
                # Configurar progresso para ocupar toda a largura
                progresso = progresso_placeholder.progress(0, text="")
                status = status_placeholder
                
                atualizar_progresso(progresso, status, "🤖 Carregando modelo Whisper", 5)
                modelo = whisper.load_model(modelo_escolhido)
                atualizar_progresso(progresso, status, "🎙️ Transcrevendo áudio", 15)
                progress_queue = queue.Queue()
                if st.session_state['chunk_processing']:
                    processar_audio_chunk(modelo, audio_path, idioma_escolhido, progress_queue)
                    result_type, resultado = progress_queue.get_nowait()
                    if result_type == "erro":
                        raise Exception(resultado)
                else:
                    resultado = modelo.transcribe(audio_path, language=None if idioma_escolhido=="auto" else idioma_escolhido)
                atualizar_progresso(progresso, status, "✅ Transcrição concluída", 50)
                del modelo
                
                # 🔊 Diarização condicional
                diarization = None
                if st.session_state['habilitar_diarizacao']:
                    atualizar_progresso(progresso, status, "👥 Inicializando diarização", 55)
                    processar_diarizacao(audio_path, HUGGINGFACE_TOKEN, progress_queue)
                    result_type, diarization = progress_queue.get_nowait()
                    if result_type == "erro":
                        raise Exception(resultado)
                    atualizar_progresso(progresso, status, "✅ Diarização concluída", 75)
                else:
                    atualizar_progresso(progresso, status, "⏩ Diarização desabilitada", 75)
                
                # 📝 Processar segmentos
                falas = []
                mapa_locutores = {}
                contador_locutor = 1
                total_segmentos = len(resultado["segments"])
                for i, segmento in enumerate(resultado["segments"]):
                    start = segmento["start"]
                    end = segmento["end"]
                    texto = f'"{segmento["text"].strip()}"'
                    speaker = "Locutor 1"  # Valor padrão quando diarização está desabilitada
                    
                    if st.session_state['habilitar_diarizacao'] and diarization:
                        speaker = "Desconhecido"
                        for turno in diarization.itertracks(yield_label=True):
                            seg_dia = turno[0]
                            locutor_original = turno[-1]
                            if seg_dia.start <= start <= seg_dia.end:
                                speaker = locutor_original
                                break
                        if speaker not in mapa_locutores and speaker != "Desconhecido":
                            mapa_locutores[speaker] = f"Locutor {contador_locutor}"
                            contador_locutor += 1
                        speaker = mapa_locutores.get(speaker, speaker)
                    
                    falas.append({
                        "tempo": f"{formatar_tempo(start)} - {formatar_tempo(end)}",
                        "locutor": speaker,
                        "texto": texto
                    })
                    if i % max(1, total_segmentos // 10) == 0:
                        progress_val = 80 + int(10*(i+1)/total_segmentos)
                        atualizar_progresso(progresso, status, "🧩 Processando segmentos", progress_val)
                
                # 📄 Gerar documentos
                doc = Document()
                doc.add_heading(f'Tabela 1 - Transcrição do Áudio "{audio_file.name}".', level=1)
                tabela = doc.add_table(rows=1, cols=3)
                tabela.style = 'Table Grid'
                hdr_cells = tabela.rows[0].cells
                hdr_cells[0].text = 'Tempo'
                hdr_cells[1].text = 'Locutor'
                hdr_cells[2].text = 'Transcrição'
                for fala in falas:
                    row_cells = tabela.add_row().cells
                    row_cells[0].text = fala["tempo"]
                    row_cells[1].text = fala["locutor"]
                    row_cells[2].text = fala["texto"]
                doc_stream = BytesIO()
                doc.save(doc_stream)
                doc_stream.seek(0)
                st.session_state["doc_word"] = doc_stream
                st.session_state["tabela_falas"] = pd.DataFrame([{"Tempo": f["tempo"], "Locutor": f["locutor"], "Transcrição": f["texto"]} for f in falas])
                st.session_state["tabela_falas"].index += 1
                st.session_state["csv_data"] = st.session_state["tabela_falas"].to_csv(index=False)
                
                # 🎬 Gerar arquivo SRT
                srt_content = criar_srt(falas)
                st.session_state["srt_data"] = srt_content
                
                atualizar_progresso(progresso, status, "✅ Processamento concluído!", 100)

        except Exception as e:
            st.error(f"⚠️ Erro durante o processamento:")
            st.code(str(e))
            st.expander("🐞 Detalhes técnicos").code(traceback.format_exc())
        finally:
            if os.path.exists(audio_path):
                os.remove(audio_path)

# 📊 Exibir resultados
if "tabela_falas" in st.session_state:
    col_btn1, col_btn2, col_btn3, col_btn4 = st.columns(4)
    with col_btn1:
        st.download_button(
            label="📄 Baixar Word",
            data=st.session_state["doc_word"],
            file_name=f"{st.session_state['audio_processado'].split('.')[0]}-{modelo_escolhido}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    with col_btn2:
        st.download_button(
            label="📊 Baixar CSV",
            data=st.session_state["csv_data"],
            file_name=f"{st.session_state['audio_processado'].split('.')[0]}-{modelo_escolhido}.csv",
            mime="text/csv"
        )
    with col_btn3:
        st.download_button(
            label="📝 Baixar SRT",
            data=st.session_state["srt_data"],
            file_name=f"{st.session_state['audio_processado'].split('.')[0]}-{modelo_escolhido}.srt",
            mime="text/plain"
        )
    with col_btn4:
        if st.button("🧹 Limpar Resultados", type="primary"):
            for key in ["tabela_falas", "doc_word", "csv_data", "srt_data", "audio_processado"]:
                st.session_state.pop(key, None)
            st.rerun()
    
    st.subheader("🎼 Transcrição Completa")
    st.dataframe(st.session_state["tabela_falas"], use_container_width=True, height=400)
    with st.expander("📈 Estatísticas Detalhadas"):
        df = st.session_state["tabela_falas"]
        col_stat1, col_stat2 = st.columns(2)
        with col_stat1:
            st.metric("Total de Falas", len(df))
            st.metric("Locutores Únicos", df['Locutor'].nunique())
        with col_stat2:
            distribuicao = df['Locutor'].value_counts()
            st.write("**Distribuição de Falas por Locutor:**")
            for locutor, count in distribuicao.items():
                st.write(f"• {locutor}: {count} falas")

# Executa a aplicação
if __name__ == "__main__":
    st.write("")  # Necessário para inicialização do Streamlit
